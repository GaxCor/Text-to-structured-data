"""
PlanetaFiscal - Procesador de Texto a Datos Estructurados
=========================================================
Convierte texto desordenado (correos, facturas, quejas) en JSON estructurado
listo para inserción en base de datos SQL, utilizando la API de OpenAI.
"""

import json
import os
import sys
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import openpyxl
import pdfplumber
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

# ── Configuración ────────────────────────────────────────────────────────────

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = BASE_DIR / "datos_entrada"
OUTPUT_DIR = BASE_DIR / "datos_salida"

MAX_REINTENTOS = 3
MODELO_OPENAI = "gpt-4o-mini"

CAMPOS_REQUERIDOS = {"nombre_cliente", "monto", "fecha", "tipo_solicitud"}
TIPOS_VALIDOS = {"Venta", "Queja", "Factura"}
EXTENSIONES_SOPORTADAS = {".txt", ".pdf", ".docx", ".xlsx", ".xls"}

# ── Logging ──────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("PlanetaFiscal")

# ── Prompt del sistema ───────────────────────────────────────────────────────

SYSTEM_PROMPT = """Eres un asistente experto en extracción de datos estructurados.
Tu tarea es analizar texto desordenado (correos, facturas, quejas, solicitudes)
y devolver EXCLUSIVAMENTE un objeto JSON válido con la siguiente estructura:

{
  "nombre_cliente": "string",
  "monto": number | null,
  "fecha": "YYYY-MM-DD",
  "tipo_solicitud": "Venta | Queja | Factura"
}

Reglas estrictas:
1. "nombre_cliente": Nombre de la persona o empresa que envía/solicita. Siempre string.
2. "monto": Monto numérico principal mencionado (sin símbolos de moneda). Si no existe, usar null.
3. "fecha": Fecha más relevante del documento en formato YYYY-MM-DD. Si hay varias, usar la principal.
4. "tipo_solicitud": SOLO puede ser "Venta", "Queja" o "Factura". Clasifica según el contenido.

IMPORTANTE:
- Responde ÚNICAMENTE con el JSON. Sin texto adicional, sin explicaciones, sin markdown.
- No envuelvas el JSON en bloques de código.
- El JSON debe ser válido y parseable directamente.
"""


# ── Funciones principales ────────────────────────────────────────────────────


def crear_cliente_openai() -> OpenAI:
    """Crea y retorna un cliente de OpenAI validando que la API key exista."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or api_key == "sk-tu-api-key-aqui":
        raise EnvironmentError(
            "OPENAI_API_KEY no configurada. "
            "Edita el archivo .env con tu API key válida."
        )
    return OpenAI(api_key=api_key)


def leer_archivo(ruta: Path) -> str:
    """
    Lee el contenido de un archivo según su extensión.
    Soporta: .txt, .pdf, .docx, .xlsx, .xls
    """
    ext = ruta.suffix.lower()

    if ext == ".txt":
        return _leer_txt(ruta)
    elif ext == ".pdf":
        return _leer_pdf(ruta)
    elif ext == ".docx":
        return _leer_docx(ruta)
    elif ext in (".xlsx", ".xls"):
        return _leer_excel(ruta)
    else:
        raise ValueError(f"Formato no soportado: {ext} (archivo: {ruta.name})")


def _leer_txt(ruta: Path) -> str:
    """Lee un archivo de texto plano con detección de encoding."""
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            return ruta.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"No se pudo decodificar el archivo: {ruta}")


def _leer_pdf(ruta: Path) -> str:
    """Extrae todo el texto de un archivo PDF."""
    texto_paginas = []
    with pdfplumber.open(ruta) as pdf:
        for pagina in pdf.pages:
            texto = pagina.extract_text()
            if texto:
                texto_paginas.append(texto)
    contenido = "\n".join(texto_paginas)
    if not contenido.strip():
        raise ValueError(f"El PDF no contiene texto extraíble: {ruta.name}")
    return contenido


def _leer_docx(ruta: Path) -> str:
    """Extrae todo el texto de un archivo Word (.docx)."""
    doc = Document(ruta)
    parrafos = [p.text for p in doc.paragraphs if p.text.strip()]

    # También extraer texto de tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                parrafos.append(" | ".join(celdas))

    contenido = "\n".join(parrafos)
    if not contenido.strip():
        raise ValueError(f"El documento Word está vacío: {ruta.name}")
    return contenido


def _leer_excel(ruta: Path) -> str:
    """Extrae todo el texto de un archivo Excel (.xlsx/.xls)."""
    wb = openpyxl.load_workbook(ruta, read_only=True, data_only=True)
    lineas = []

    for hoja in wb.sheetnames:
        ws = wb[hoja]
        lineas.append(f"--- Hoja: {hoja} ---")
        for fila in ws.iter_rows(values_only=True):
            valores = [str(v) if v is not None else "" for v in fila]
            linea = " | ".join(valores)
            if linea.strip(" |"):
                lineas.append(linea)

    wb.close()
    contenido = "\n".join(lineas)
    if not contenido.strip():
        raise ValueError(f"El archivo Excel está vacío: {ruta.name}")
    return contenido


def validar_json(respuesta_texto: str) -> dict:
    """
    Valida que la respuesta sea un JSON válido con todos los campos requeridos
    y tipos correctos. Lanza ValueError si la validación falla.
    """
    # Limpiar posibles envolturas de markdown
    texto = respuesta_texto.strip()
    if texto.startswith("```"):
        lineas = texto.split("\n")
        lineas = [l for l in lineas if not l.strip().startswith("```")]
        texto = "\n".join(lineas).strip()

    # Parsear JSON
    try:
        datos = json.loads(texto)
    except json.JSONDecodeError as e:
        raise ValueError(f"Respuesta no es JSON válido: {e}")

    # Si la IA devuelve una lista (múltiples registros), tomar el primero
    if isinstance(datos, list):
        if len(datos) == 0:
            raise ValueError("La IA devolvió una lista vacía.")
        datos = datos[0]

    # Verificar que sea un diccionario
    if not isinstance(datos, dict):
        raise ValueError(f"Se esperaba un objeto JSON, se recibió: {type(datos).__name__}")

    # Verificar campos requeridos
    campos_faltantes = CAMPOS_REQUERIDOS - set(datos.keys())
    if campos_faltantes:
        raise ValueError(f"Campos faltantes en el JSON: {campos_faltantes}")

    # Validar tipos
    if not isinstance(datos["nombre_cliente"], str) or not datos["nombre_cliente"].strip():
        raise ValueError("'nombre_cliente' debe ser un string no vacío.")

    if datos["monto"] is not None:
        if not isinstance(datos["monto"], (int, float)):
            raise ValueError(f"'monto' debe ser numérico o null, se recibió: {type(datos['monto']).__name__}")

    if not isinstance(datos["fecha"], str):
        raise ValueError("'fecha' debe ser un string en formato YYYY-MM-DD.")
    try:
        datetime.strptime(datos["fecha"], "%Y-%m-%d")
    except ValueError:
        raise ValueError(f"'fecha' no tiene formato YYYY-MM-DD válido: {datos['fecha']}")

    if datos["tipo_solicitud"] not in TIPOS_VALIDOS:
        raise ValueError(
            f"'tipo_solicitud' inválido: '{datos['tipo_solicitud']}'. "
            f"Valores permitidos: {TIPOS_VALIDOS}"
        )

    return datos


def extraer_datos(cliente: OpenAI, texto: str) -> dict:
    """
    Envía el texto a OpenAI y extrae datos estructurados.
    Reintenta hasta MAX_REINTENTOS veces si la respuesta es inválida.
    """
    for intento in range(1, MAX_REINTENTOS + 1):
        logger.info(f"  Intento {intento}/{MAX_REINTENTOS}...")

        try:
            respuesta = cliente.chat.completions.create(
                model=MODELO_OPENAI,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"Extrae los datos del siguiente texto:\n\n{texto}"},
                ],
            )

            texto_respuesta = respuesta.choices[0].message.content
            logger.debug(f"  Respuesta cruda: {texto_respuesta}")

            datos = validar_json(texto_respuesta)
            logger.info(f"  JSON válido extraído exitosamente.")
            return datos

        except ValueError as e:
            logger.warning(f"  Validación fallida: {e}")
            if intento == MAX_REINTENTOS:
                raise RuntimeError(
                    f"Falló la extracción después de {MAX_REINTENTOS} intentos. "
                    f"Último error: {e}"
                )

        except Exception as e:
            logger.error(f"  Error de API: {e}")
            if intento == MAX_REINTENTOS:
                raise RuntimeError(
                    f"Error de API después de {MAX_REINTENTOS} intentos: {e}"
                )

    # Nunca debería llegar aquí
    raise RuntimeError("Error inesperado en extraer_datos.")


def procesar_archivos() -> list[dict]:
    """
    Procesa todos los archivos soportados del directorio de entrada.
    Formatos: .txt, .pdf, .docx, .xlsx, .xls
    Retorna una lista de resultados estructurados.
    """
    cliente = crear_cliente_openai()

    # Crear directorio de salida
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Buscar archivos de todos los formatos soportados
    archivos = sorted(
        f for f in INPUT_DIR.iterdir()
        if f.is_file() and f.suffix.lower() in EXTENSIONES_SOPORTADAS
    )
    if not archivos:
        logger.warning(
            f"No se encontraron archivos soportados en {INPUT_DIR}. "
            f"Extensiones válidas: {', '.join(EXTENSIONES_SOPORTADAS)}"
        )
        return []

    logger.info(f"Se encontraron {len(archivos)} archivo(s) para procesar.\n")

    resultados = []
    errores = []

    for archivo in archivos:
        logger.info(f"Procesando: {archivo.name}")

        try:
            texto = leer_archivo(archivo)
            datos = extraer_datos(cliente, texto)

            # Agregar metadatos
            resultado = {
                "archivo_origen": archivo.name,
                "datos_extraidos": datos,
                "procesado_en": datetime.now().isoformat(),
            }
            resultados.append(resultado)

            # Guardar resultado individual
            salida_individual = OUTPUT_DIR / f"{archivo.stem}_resultado.json"
            salida_individual.write_text(
                json.dumps(datos, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            logger.info(f"  Guardado: {salida_individual.name}\n")

        except RuntimeError as e:
            logger.error(f"  ERROR FATAL: {e}\n")
            errores.append({"archivo": archivo.name, "error": str(e)})

    # Guardar resumen consolidado
    resumen = {
        "total_procesados": len(resultados),
        "total_errores": len(errores),
        "fecha_procesamiento": datetime.now().isoformat(),
        "resultados": resultados,
        "errores": errores,
    }

    ruta_resumen = OUTPUT_DIR / "resumen_completo.json"
    ruta_resumen.write_text(
        json.dumps(resumen, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    logger.info(f"Resumen guardado en: {ruta_resumen}")

    # Resumen final
    logger.info(f"\n{'='*50}")
    logger.info(f"RESUMEN: {len(resultados)} exitosos, {len(errores)} errores")
    logger.info(f"{'='*50}")

    return resultados


# ── Punto de entrada ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    try:
        resultados = procesar_archivos()
        if resultados:
            print("\n📋 Ejemplo de salida JSON extraída:")
            print(json.dumps(resultados[0]["datos_extraidos"], indent=2, ensure_ascii=False))
    except EnvironmentError as e:
        logger.error(e)
        sys.exit(1)
    except KeyboardInterrupt:
        logger.info("\nProceso interrumpido por el usuario.")
        sys.exit(0)
